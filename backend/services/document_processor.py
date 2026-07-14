"""
Document processor service for parsing PDF, DOCX, and TXT files.
Extracts text and metadata from various document formats.
"""
import io
import uuid
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
import pypdf
from docx import Document as DocxDocument
from config import settings
from config.logging_config import get_logger

logger = get_logger(__name__)


class DocumentProcessorService:
    """Service for processing various document formats."""

    def __init__(self):
        """Initialize the document processor service."""
        self.supported_formats = {'.pdf', '.docx', '.txt', '.md', '.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
        self._image_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.bmp'}
        logger.info("Document processor service initialized")

    def is_supported(self, filename: str) -> bool:
        """Check if file format is supported."""
        ext = Path(filename).suffix.lower()
        return ext in self.supported_formats

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        extract_images: bool = True,
    ) -> Dict[str, Any]:
        """
        Process a document and extract text and metadata.

        Args:
            file_content: Document file bytes
            filename: Name of the file
            extract_images: Whether to extract images for OCR

        Returns:
            Dictionary with extracted text, metadata, and images
        """
        document_id = str(uuid.uuid4())

        try:
            logger.info(f"Processing document: {filename}")

            file_ext = Path(filename).suffix.lower()

            if file_ext == '.pdf':
                result = await self._process_pdf(file_content, filename, extract_images)
            elif file_ext == '.docx':
                result = await self._process_docx(file_content, filename)
            elif file_ext in {'.txt', '.md'}:
                result = await self._process_text(file_content, filename)
            elif file_ext in self._image_formats:
                result = await self._process_image(file_content, filename)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

            result["document_id"] = document_id
            result["filename"] = filename
            result["file_type"] = file_ext[1:]  # Remove dot
            result["file_size"] = len(file_content)

            logger.info(
                f"Document processed: {len(result.get('text', ''))} chars extracted, "
                f"{len(result.get('images', []))} images extracted"
            )

            return result

        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise

    async def _process_pdf(
        self,
        content: bytes,
        filename: str,
        extract_images: bool = True,
    ) -> Dict[str, Any]:
        """Process PDF document with OCR fallback for scanned documents."""
        text_parts = []
        images = []
        metadata = {}
        extraction_method = "pypdf"

        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = pypdf.PdfReader(pdf_file)

            # Get metadata
            if pdf_reader.metadata:
                metadata = {
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", ""),
                    "producer": pdf_reader.metadata.get("/Producer", ""),
                    "pages": len(pdf_reader.pages),
                }

            page_count = len(pdf_reader.pages) or 1

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append({
                            "page": page_num + 1,
                            "text": page_text.strip(),
                        })
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")

            # Calculate character density for OCR fallback decision
            total_chars = sum(len(part["text"]) for part in text_parts)
            avg_chars_per_page = total_chars / page_count

            # OCR fallback if text extraction yields low results
            if (avg_chars_per_page < settings.ocr_min_chars_per_page
                    and settings.ocr_fallback_enabled):
                logger.info(
                    f"Low text density ({avg_chars_per_page:.0f} chars/page) in {filename}, "
                    f"attempting OCR fallback"
                )
                ocr_result = await self._ocr_fallback(content, filename)
                if ocr_result:
                    return ocr_result

            # Combine all text from pypdf extraction
            full_text = "\n\n".join(
                part["text"] for part in text_parts
            )

            return {
                "text": full_text,
                "pages": text_parts,
                "images": images,
                "metadata": metadata,
                "extraction_method": extraction_method,
            }

        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

    async def _ocr_fallback(
        self,
        content: bytes,
        filename: str,
    ) -> Optional[Dict[str, Any]]:
        """Attempt OCR extraction as fallback for scanned PDFs."""
        try:
            from services.ocr_service import get_ocr_service
            ocr_service = get_ocr_service()

            if ocr_service.provider == "none":
                logger.warning("OCR fallback requested but no OCR provider configured")
                return None

            ocr_result = await ocr_service.extract_with_structure(content)

            if not ocr_result.get("text"):
                logger.warning("OCR extraction returned no text")
                return None

            # Convert OCR result to standard format
            pages = []
            for page_info in ocr_result.get("pages", []):
                pages.append({
                    "page": page_info.get("page_number", 1),
                    "text": page_info.get("text", ""),
                })

            logger.info(f"OCR fallback successful for {filename}: {len(ocr_result['text'])} chars")

            return {
                "text": ocr_result["text"],
                "pages": pages or [{"page": 1, "text": ocr_result["text"]}],
                "images": [],
                "metadata": {},
                "tables": ocr_result.get("tables", []),
                "extraction_method": "ocr",
            }

        except Exception as e:
            logger.error(f"OCR fallback failed: {e}")
            return None

    async def _process_docx(
        self,
        content: bytes,
        filename: str,
    ) -> Dict[str, Any]:
        """Process DOCX document."""
        try:
            doc_file = io.BytesIO(content)
            doc = DocxDocument(doc_file)

            # Extract text from paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())

            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)

            # Extract metadata
            metadata = {}
            if doc.core_properties.title:
                metadata["title"] = doc.core_properties.title
            if doc.core_properties.author:
                metadata["author"] = doc.core_properties.author

            # Combine text
            full_text = "\n\n".join(text_parts)

            return {
                "text": full_text,
                "pages": [{"page": 1, "text": full_text}],
                "images": [],
                "tables": tables,
                "metadata": metadata,
            }

        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise

    async def _process_text(
        self,
        content: bytes,
        filename: str,
    ) -> Dict[str, Any]:
        """Process plain text file."""
        try:
            # Try UTF-8 first
            try:
                text = content.decode('utf-8')
            except UnicodeDecodeError:
                # Try other encodings
                for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        text = content.decode(encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    text = content.decode('utf-8', errors='ignore')

            return {
                "text": text.strip(),
                "pages": [{"page": 1, "text": text.strip()}],
                "images": [],
                "metadata": {"encoding": "utf-8"},
                "extraction_method": "text",
            }

        except Exception as e:
            logger.error(f"Error processing text file: {e}")
            raise

    async def _process_image(
        self,
        content: bytes,
        filename: str,
    ) -> Dict[str, Any]:
        """Process image file using OCR."""
        try:
            from services.ocr_service import get_ocr_service
            ocr_service = get_ocr_service()

            if ocr_service.provider == "none":
                raise ValueError(
                    "OCR not configured. Set AZURE_DOC_INTELLIGENCE_ENDPOINT and "
                    "AZURE_DOC_INTELLIGENCE_KEY to process images."
                )

            logger.info(f"Processing image with OCR: {filename}")
            ocr_result = await ocr_service.extract_text(content, filename)

            if ocr_result.get("error"):
                raise ValueError(f"OCR failed: {ocr_result['error']}")

            text = ocr_result.get("text", "")

            return {
                "text": text,
                "pages": [{"page": 1, "text": text}],
                "images": [],
                "metadata": {
                    "ocr_provider": ocr_result.get("provider", "unknown"),
                    "line_count": ocr_result.get("line_count", 0),
                },
                "extraction_method": "ocr",
            }

        except Exception as e:
            logger.error(f"Error processing image: {e}")
            raise

    async def extract_text_for_indexing(
        self,
        file_content: bytes,
        filename: str,
    ) -> str:
        """
        Extract just the text from a document for indexing.
        Lightweight method for bulk processing.
        """
        result = await self.process_document(file_content, filename, extract_images=False)
        return result.get("text", "")
